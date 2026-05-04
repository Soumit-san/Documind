"""
Export Service — Generates PDF and DOCX files from analysis data.
Uses ReportLab for PDF and python-docx for DOCX generation.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger("documind.export")


def generate_pdf(
    title: str,
    executive_summary: str = "",
    sections: Optional[list[dict]] = None,
    entities: Optional[dict] = None,
    citations: Optional[list[dict]] = None,
    chat_messages: Optional[list[dict]] = None,
) -> bytes:
    """
    Generate a styled PDF document from analysis results.
    Returns the PDF as bytes.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=25 * mm, rightMargin=25 * mm,
        topMargin=25 * mm, bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    primary = HexColor("#A8FF00")
    dark_bg = HexColor("#1A1A1A")

    # Custom styles
    styles.add(ParagraphStyle(
        "DocuTitle", parent=styles["Title"],
        fontSize=22, leading=28, textColor=HexColor("#333333"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "DocuSubtitle", parent=styles["Normal"],
        fontSize=10, textColor=HexColor("#888888"),
        spaceAfter=20,
    ))
    styles.add(ParagraphStyle(
        "SectionHeading", parent=styles["Heading2"],
        fontSize=14, leading=18, textColor=HexColor("#222222"),
        spaceBefore=16, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "BodyText14", parent=styles["Normal"],
        fontSize=11, leading=17, textColor=HexColor("#333333"),
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        "ChatUser", parent=styles["Normal"],
        fontSize=10, leading=14, textColor=HexColor("#1A1A1A"),
        leftIndent=20, spaceBefore=6, spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        "ChatAssistant", parent=styles["Normal"],
        fontSize=10, leading=14, textColor=HexColor("#333333"),
        leftIndent=20, spaceBefore=2, spaceAfter=8,
    ))

    elements = []

    # --- Header ---
    elements.append(Paragraph("DOCUMIND AI", styles["DocuTitle"]))
    elements.append(Paragraph(f"Analysis Report — {title}", styles["DocuSubtitle"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=HexColor("#CCCCCC"), spaceAfter=16))

    # --- Executive Summary ---
    if executive_summary:
        elements.append(Paragraph("Executive Summary", styles["SectionHeading"]))
        elements.append(Paragraph(executive_summary, styles["BodyText14"]))
        elements.append(Spacer(1, 10))

    # --- Section Breakdown ---
    if sections:
        elements.append(Paragraph("Section Breakdown", styles["SectionHeading"]))
        for sec in sections:
            sec_title = sec.get("title", "Untitled Section")
            sec_summary = sec.get("summary", "")
            page = sec.get("page")
            label = f"<b>{sec_title}</b>"
            if page:
                label += f" <i>(Page {page})</i>"
            elements.append(Paragraph(label, styles["BodyText14"]))
            elements.append(Paragraph(sec_summary, styles["BodyText14"]))
        elements.append(Spacer(1, 10))

    # --- Entities ---
    if entities:
        elements.append(Paragraph("Key Entities", styles["SectionHeading"]))
        for category, items in entities.items():
            if items:
                label = category.replace("_", " ").title()
                elements.append(Paragraph(
                    f"<b>{label}:</b> {', '.join(items)}",
                    styles["BodyText14"],
                ))
        elements.append(Spacer(1, 10))

    # --- Citations ---
    if citations:
        elements.append(Paragraph("Citations", styles["SectionHeading"]))
        for i, c in enumerate(citations):
            cite_text = c.get("text", "")[:200]
            page = c.get("page")
            section = c.get("section", "")
            ref = f"[{i + 1}] "
            if section:
                ref += f"{section} "
            if page:
                ref += f"(Page {page}) "
            ref += f"— \"{cite_text}...\""
            elements.append(Paragraph(ref, styles["BodyText14"]))
        elements.append(Spacer(1, 10))

    # --- Chat Transcript ---
    if chat_messages:
        elements.append(HRFlowable(width="100%", thickness=1, color=HexColor("#DDDDDD"), spaceBefore=12, spaceAfter=12))
        elements.append(Paragraph("Chat Transcript", styles["SectionHeading"]))
        for msg in chat_messages:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            if role == "user":
                elements.append(Paragraph(f"<b>You:</b> {content}", styles["ChatUser"]))
            else:
                elements.append(Paragraph(f"<b>DocuMind AI:</b> {content}", styles["ChatAssistant"]))

    # --- Footer ---
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#DDDDDD"), spaceAfter=6))
    elements.append(Paragraph(
        "Generated by DocuMind AI — Intelligent Document Assistant",
        ParagraphStyle("Footer", parent=styles["Normal"], fontSize=8, textColor=HexColor("#AAAAAA"), alignment=TA_CENTER),
    ))

    doc.build(elements)
    pdf_bytes = buf.getvalue()
    buf.close()
    logger.info("Generated PDF: %d bytes", len(pdf_bytes))
    return pdf_bytes


def generate_docx(
    title: str,
    executive_summary: str = "",
    sections: Optional[list[dict]] = None,
    entities: Optional[dict] = None,
    citations: Optional[list[dict]] = None,
    chat_messages: Optional[list[dict]] = None,
) -> bytes:
    """
    Generate a styled DOCX document from analysis results.
    Returns the DOCX as bytes.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Style tweaks ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Title ---
    title_para = doc.add_heading("DOCUMIND AI", level=0)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    subtitle = doc.add_paragraph(f"Analysis Report — {title}")
    subtitle.style = doc.styles["Normal"]
    for run in subtitle.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")  # spacer

    # --- Executive Summary ---
    if executive_summary:
        doc.add_heading("⚡ Executive Summary", level=1)
        doc.add_paragraph(executive_summary)

    # --- Section Breakdown ---
    if sections:
        doc.add_heading("📑 Section Breakdown", level=1)
        for sec in sections:
            sec_title = sec.get("title", "Untitled Section")
            sec_summary = sec.get("summary", "")
            page = sec.get("page")
            heading_text = sec_title
            if page:
                heading_text += f" (Page {page})"
            doc.add_heading(heading_text, level=2)
            doc.add_paragraph(sec_summary)

    # --- Entities ---
    if entities:
        doc.add_heading("🏷 Key Entities", level=1)
        for category, items in entities.items():
            if items:
                label = category.replace("_", " ").title()
                para = doc.add_paragraph()
                run_bold = para.add_run(f"{label}: ")
                run_bold.bold = True
                para.add_run(", ".join(items))

    # --- Citations ---
    if citations:
        doc.add_heading("📎 Citations", level=1)
        for i, c in enumerate(citations):
            cite_text = c.get("text", "")[:200]
            page = c.get("page")
            section = c.get("section", "")
            ref = f"[{i + 1}] "
            if section:
                ref += f"{section} "
            if page:
                ref += f"(Page {page}) "
            ref += f'— "{cite_text}..."'
            doc.add_paragraph(ref)

    # --- Chat Transcript ---
    if chat_messages:
        doc.add_heading("💬 Chat Transcript", level=1)
        for msg in chat_messages:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            para = doc.add_paragraph()
            role_label = "You" if role == "user" else "DocuMind AI"
            run_bold = para.add_run(f"{role_label}: ")
            run_bold.bold = True
            para.add_run(content)

    # --- Footer ---
    doc.add_paragraph("")
    footer_para = doc.add_paragraph("Generated by DocuMind AI — Intelligent Document Assistant")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    buf.close()
    logger.info("Generated DOCX: %d bytes", len(docx_bytes))
    return docx_bytes
